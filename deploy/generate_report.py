"""Génère le rapport .docx de déploiement GlobalAfrica+."""
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime

# ── Couleurs de marque ────────────────────────────────────
GOLD   = RGBColor(0xC5, 0xA5, 0x4E)
RED    = RGBColor(0xC4, 0x34, 0x4E)
BLACK  = RGBColor(0x1A, 0x1A, 0x1A)
GREY   = RGBColor(0x55, 0x55, 0x55)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
GREEN  = RGBColor(0x2E, 0x7D, 0x32)

doc = Document()

# Marges
for s in doc.sections:
    s.top_margin = Cm(2)
    s.bottom_margin = Cm(2)
    s.left_margin = Cm(2.2)
    s.right_margin = Cm(2.2)

# Style de base
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

def set_cell_bg(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)

def h1(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(20)
    r.font.color.rgb = GOLD
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)

def h2(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(15)
    r.font.color.rgb = BLACK
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)

def h3(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(12)
    r.font.color.rgb = RED
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)

def para(text, bold=False, color=None, italic=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    if color:
        r.font.color.rgb = color
    return p

def code(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = 'Consolas'
    r.font.size = Pt(9.5)
    p.paragraph_format.left_indent = Cm(0.4)
    p.paragraph_format.space_after = Pt(4)
    # fond gris clair
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), 'F2F2F2')
    p_pr.append(shd)

def bullet(text):
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(text)

def table_2col(rows, col1_w=5, col2_w=11, header=None):
    t = doc.add_table(rows=1 if header else 0, cols=2)
    t.style = 'Light Grid Accent 1'
    if header:
        hr = t.rows[0].cells
        hr[0].text = header[0]
        hr[1].text = header[1]
        for c in hr:
            set_cell_bg(c, '1A1A1A')
            for p in c.paragraphs:
                for r in p.runs:
                    r.font.color.rgb = WHITE
                    r.bold = True
    for k, v in rows:
        row = t.add_row().cells
        row[0].text = k
        row[1].text = v
        for p in row[0].paragraphs:
            for r in p.runs:
                r.bold = True
    t.columns[0].width = Cm(col1_w)
    t.columns[1].width = Cm(col2_w)

def hr_line():
    p = doc.add_paragraph()
    p_pr = p._p.get_or_add_pPr()
    p_bdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:color'), 'C5A54E')
    p_bdr.append(bottom)
    p_pr.append(p_bdr)

# ═══════════════════════════════════════════════════════════
# PAGE DE GARDE
# ═══════════════════════════════════════════════════════════
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run('GlobalAfrica+')
r.bold = True
r.font.size = Pt(38)
r.font.color.rgb = GOLD
t.paragraph_format.space_before = Pt(120)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("Rapport de mise en production")
r.font.size = Pt(22)
r.font.color.rgb = BLACK
r.bold = True

sub2 = doc.add_paragraph()
sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub2.add_run("Pipeline CI/CD et déploiement VPS Hostinger")
r.font.size = Pt(14)
r.font.color.rgb = RED
r.italic = True

doc.add_paragraph().paragraph_format.space_before = Pt(60)

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
for line in [
    ("Domaine : ", "https://globalafricaplus.com"),
    ("Repository : ", "github.com/oksam90/globalafricaplus-web"),
    ("VPS : ", "srv1535932 (Hostinger)"),
    ("Date de mise en ligne : ", "18 avril 2026"),
]:
    r1 = info.add_run(line[0])
    r1.bold = True
    r2 = info.add_run(line[1] + "\n")
    r2.font.color.rgb = GREY

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# 1. SOMMAIRE
# ═══════════════════════════════════════════════════════════
h1("Sommaire")
sommaire = [
    "1. Contexte et objectifs",
    "2. Architecture technique",
    "3. Historique des commits",
    "4. Infrastructure VPS – provisioning",
    "5. Configuration SSH et clés de déploiement",
    "6. Configuration DNS",
    "7. SSL / HTTPS (Let's Encrypt)",
    "8. Configuration de l'environnement Laravel (.env)",
    "9. Pipeline CI/CD GitHub Actions",
    "10. Script de déploiement atomique",
    "11. Exécution des déploiements",
    "12. Validation finale et statut",
    "13. Procédures post-déploiement",
]
for s in sommaire:
    p = doc.add_paragraph()
    p.add_run(s).font.size = Pt(11)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# 1. CONTEXTE
# ═══════════════════════════════════════════════════════════
h1("1. Contexte et objectifs")

para("GlobalAfrica+ est une plateforme panafricaine (Laravel 11 + Vue 3 SPA) connectant entrepreneurs africains, investisseurs, gouvernements, chercheurs d'emploi et mentors avec la diaspora.")

para("Objectifs de cette mise en production :")
bullet("Mettre en place une pipeline CI/CD entièrement automatisée (push main → déploiement)")
bullet("Déployer l'application sur un VPS Hostinger dédié")
bullet("Configurer le domaine globalafricaplus.com avec HTTPS (Let's Encrypt)")
bullet("Assurer zero downtime via un déploiement atomique (releases + symlink)")
bullet("Sécuriser l'environnement de production (firewall, fail2ban, HSTS, secrets)")

# ═══════════════════════════════════════════════════════════
# 2. ARCHITECTURE
# ═══════════════════════════════════════════════════════════
h1("2. Architecture technique")

h2("2.1 Stack applicatif")
table_2col([
    ("Framework backend", "Laravel 11 (PHP 8.2)"),
    ("Frontend", "Vue 3 + Pinia + Vite"),
    ("UI / Styling", "Tailwind CSS v4 (dark mode class-based)"),
    ("Charte graphique", "Or #C5A54E · Noir #1A1A1A · Rouge cerise #C4344E"),
    ("Auth", "Laravel Sanctum (session + cookie)"),
    ("KYC", "IDNorm (mode sandbox au démarrage)"),
], header=("Composant", "Technologie"))

h2("2.2 Stack serveur (VPS)")
table_2col([
    ("OS", "Ubuntu 22.04 LTS"),
    ("Serveur web", "Nginx 1.24"),
    ("PHP", "8.2-FPM (opcache activé)"),
    ("Base de données", "MySQL 8"),
    ("Cache / Queue / Session", "Redis"),
    ("Workers", "Supervisor (2 process queue:work)"),
    ("SSL", "Let's Encrypt (certbot auto-renew)"),
    ("Firewall", "UFW + fail2ban"),
], header=("Composant", "Technologie"))

h2("2.3 Arborescence de déploiement")
code("""/var/www/globalafricaplus.com/
├── current         → symlink vers la release active
├── releases/
│   ├── 20260418071200_f3a08870/
│   ├── 20260418071427_00d5d84a/
│   └── …           (5 dernières conservées)
└── shared/
    ├── .env        (secrets de production)
    ├── storage/    (uploads, logs, cache)
    └── bootstrap/cache/""")

para("Le déploiement est atomique : un mv -Tf sur le symlink current bascule instantanément vers la nouvelle release. Zéro interruption de service.", italic=True, color=GREY)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# 3. COMMITS
# ═══════════════════════════════════════════════════════════
h1("3. Historique des commits")

para("Les commits réalisés dans le cadre de cette mise en production :")

t = doc.add_table(rows=1, cols=3)
t.style = 'Light Grid Accent 1'
hdr = t.rows[0].cells
hdr[0].text = "SHA"
hdr[1].text = "Date"
hdr[2].text = "Message"
for c in hdr:
    set_cell_bg(c, '1A1A1A')
    for p in c.paragraphs:
        for r in p.runs:
            r.font.color.rgb = WHITE
            r.bold = True

commits = [
    ("00d5d84", "18/04/2026 07:19", "ci: verify sudoers fix + full green deploy"),
    ("f3a0887", "18/04/2026 07:12", "fix(deploy): tolerate chmod failures on pre-existing shared files"),
    ("837aed7", "18/04/2026 07:10", "ci: trigger first deployment to globalafricaplus.com"),
    ("3f43918", "18/04/2026 04:21", "ci: add GitHub Actions pipeline + Hostinger VPS deployment"),
    ("6a027e2", "18/04/2026 01:23", "all projet files"),
    ("761a8e6", "18/04/2026 01:17", "premier commit du projet globalafricaplus"),
]
for sha, date, msg in commits:
    row = t.add_row().cells
    row[0].text = sha
    row[1].text = date
    row[2].text = msg
    for p in row[0].paragraphs:
        for r in p.runs:
            r.font.name = 'Consolas'
            r.font.size = Pt(9)

# ═══════════════════════════════════════════════════════════
# 4. PROVISIONING
# ═══════════════════════════════════════════════════════════
h1("4. Infrastructure VPS – provisioning")

para("Le script deploy/provision.sh automatise l'installation complète du VPS depuis un Ubuntu 22.04 fraîchement installé.")

h3("4.1 Récupération et exécution du script")
code("""# Sur le VPS en root
cd /tmp
wget https://raw.githubusercontent.com/oksam90/globalafricaplus-web/main/deploy/provision.sh
wget https://raw.githubusercontent.com/oksam90/globalafricaplus-web/main/deploy/nginx.conf
wget https://raw.githubusercontent.com/oksam90/globalafricaplus-web/main/deploy/supervisor-worker.conf
chmod +x provision.sh
bash provision.sh""")

h3("4.2 Opérations réalisées par provision.sh")
bullet("Installation des paquets : UFW, fail2ban, PHP 8.2 + extensions, Composer, Nginx, MySQL, Redis, Node 20, Supervisor, Certbot")
bullet("Création de l'utilisateur système deploy (groupe www-data)")
bullet("Création de l'arborescence /var/www/globalafricaplus.com/{releases,shared/{storage,bootstrap/cache}}")
bullet("Création d'un template .env dans shared/")
bullet("Création de la base de données MySQL globalafricaplus + utilisateur gapuser avec mot de passe aléatoire")
bullet("Configuration de Supervisor pour 2 workers Laravel queue:work")
bullet("Configuration du cron scheduler Laravel (toutes les minutes)")
bullet("Configuration du firewall UFW (ports 22, 80, 443)")
bullet("Activation de fail2ban (jail ssh + nginx)")
bullet("Configuration sudoers : deploy peut reload php-fpm et restart supervisorctl sans mot de passe")

h3("4.3 Résultat de la vérification post-provisioning")
code("""systemctl is-active nginx php8.2-fpm mysql redis-server supervisor
→ active (x5)

id deploy
→ uid=1001(deploy) gid=1001(deploy) groupes=1001(deploy),33(www-data)

ufw status
→ Status: active
→ 22/tcp ALLOW Anywhere
→ 80/tcp ALLOW Anywhere
→ 443/tcp ALLOW Anywhere""")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# 5. SSH
# ═══════════════════════════════════════════════════════════
h1("5. Configuration SSH et clés de déploiement")

para("Une paire de clés ed25519 dédiée au CI/CD a été générée pour limiter la surface d'attaque.")

h3("5.1 Génération de la clé")
code("""ssh-keygen -t ed25519 -f ga_deploy_key -C "ci@globalafricaplus" -N ""
→ Your identification has been saved in ga_deploy_key
→ Your public key has been saved in ga_deploy_key.pub
→ SHA256:cJ1dn1AsSi7RkDJRbjJ054XxE1T0TWeyHzTYu4uIymg""")

h3("5.2 Installation de la clé publique pour l'utilisateur deploy")
code("""mkdir -p /home/deploy/.ssh
cat ~/ga_deploy_key.pub >> /home/deploy/.ssh/authorized_keys
chown -R deploy:deploy /home/deploy/.ssh
chmod 700 /home/deploy/.ssh
chmod 600 /home/deploy/.ssh/authorized_keys""")

h3("5.3 Secrets GitHub Actions")
table_2col([
    ("SSH_HOST", "IP publique du VPS Hostinger"),
    ("SSH_PORT", "22"),
    ("SSH_USER", "deploy"),
    ("SSH_PRIVATE_KEY", "Contenu du fichier ga_deploy_key (clé privée)"),
], header=("Secret", "Valeur"))

h3("5.4 Sudoers du user deploy")
code("""# /etc/sudoers.d/deploy-globalafricaplus
deploy ALL=(root) NOPASSWD: /bin/systemctl reload php8.2-fpm
deploy ALL=(root) NOPASSWD: /usr/bin/supervisorctl restart globalafricaplus-worker*""")

# ═══════════════════════════════════════════════════════════
# 6. DNS
# ═══════════════════════════════════════════════════════════
h1("6. Configuration DNS")

para("Dans le panel Hostinger, deux enregistrements A ont été créés pour pointer vers le VPS :")

table_2col([
    ("A @", "IP du VPS (TTL 3600)"),
    ("A www", "IP du VPS (TTL 3600)"),
], header=("Enregistrement", "Valeur"))

para("Vérification de la propagation :")
code("""nslookup globalafricaplus.com
nslookup www.globalafricaplus.com
→ Retournent bien l'IP du VPS""")

# ═══════════════════════════════════════════════════════════
# 7. SSL
# ═══════════════════════════════════════════════════════════
h1("7. SSL / HTTPS (Let's Encrypt)")

h3("7.1 Problème rencontré")
para("La configuration Nginx initiale référençait les certificats SSL avant leur émission, provoquant un échec de nginx -t et bloquant Certbot (problème œuf-et-poule).")

h3("7.2 Solution appliquée")
bullet("Création d'une configuration Nginx HTTP-only temporaire (port 80 uniquement)")
bullet("Exécution de Certbot pour obtenir les certificats")
bullet("Restauration de la configuration complète (HTTPS + HSTS + headers de sécurité)")

h3("7.3 Commande Certbot")
code("""certbot --nginx \\
  -d globalafricaplus.com \\
  -d www.globalafricaplus.com \\
  --non-interactive --agree-tos \\
  -m admin@globalafricaplus.com \\
  --redirect

→ Successfully deployed certificate for globalafricaplus.com
→ Successfully deployed certificate for www.globalafricaplus.com""")

h3("7.4 Renouvellement automatique")
code("""systemctl is-active certbot.timer → active
certbot renew --dry-run → Congratulations, all simulated renewals succeeded""")

h3("7.5 Tests de validation")
code("""curl -I https://globalafricaplus.com
→ HTTP/1.1 200 OK
→ Strict-Transport-Security: max-age=31536000; includeSubDomains
→ X-Frame-Options: SAMEORIGIN
→ X-Content-Type-Options: nosniff
→ Referrer-Policy: strict-origin-when-cross-origin
→ Permissions-Policy: geolocation=(), microphone=(), camera=()

curl -I http://globalafricaplus.com
→ HTTP/1.1 301 Moved Permanently
→ Location: https://globalafricaplus.com/""")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# 8. .ENV
# ═══════════════════════════════════════════════════════════
h1("8. Configuration de l'environnement Laravel")

para("Le fichier /var/www/globalafricaplus.com/shared/.env contient la configuration production, monté dans chaque release via symlink.")

h3("8.1 Configuration base de données")
code("""DB_CONNECTION=mysql
DB_HOST=127.0.0.1
DB_PORT=3306
DB_DATABASE=globalafricaplus
DB_USERNAME=gapuser
DB_PASSWORD=<généré par provision.sh>""")

h3("8.2 Redis (cache, queue, sessions)")
code("""CACHE_DRIVER=redis
QUEUE_CONNECTION=redis
SESSION_DRIVER=redis
REDIS_HOST=127.0.0.1
REDIS_PORT=6379""")

h3("8.3 Mail (mode log en attendant SMTP Hostinger)")
code("""MAIL_MAILER=log
MAIL_FROM_ADDRESS=no-reply@globalafricaplus.com
MAIL_FROM_NAME="GlobalAfrica+\"""")

h3("8.4 IDNorm KYC (sandbox au démarrage)")
code("""IDNORM_MODE=sandbox
IDNORM_BASE_URL=https://sandbox.idnorm.com/v1
IDNORM_API_KEY=placeholder
IDNORM_API_SECRET=placeholder
IDNORM_WEBHOOK_SECRET=placeholder""")

h3("8.5 Sanctum et cookies sécurisés")
code("""SANCTUM_STATEFUL_DOMAINS=globalafricaplus.com,www.globalafricaplus.com
SESSION_DOMAIN=.globalafricaplus.com
SESSION_SECURE_COOKIE=true""")

h3("8.6 Validation connexion DB")
code("""mysql -u\"$DB_USERNAME\" -p\"$DB_PASSWORD\" -e \"SELECT 'DB OK';\" \"$DB_DATABASE\"
→ DB OK""")

# ═══════════════════════════════════════════════════════════
# 9. CI/CD
# ═══════════════════════════════════════════════════════════
h1("9. Pipeline CI/CD GitHub Actions")

para("Le workflow .github/workflows/deploy.yml se déclenche automatiquement sur chaque push sur la branche main.")

h3("9.1 Job 1 — test-and-build")
bullet("Checkout du code")
bullet("Setup PHP 8.2 + extensions")
bullet("Cache Composer")
bullet("composer install --no-dev --optimize-autoloader")
bullet("PHP lint (vérification syntaxique)")
bullet("Setup Node 20 + cache npm")
bullet("npm ci && npm run build (Vite : compile Tailwind + Vue)")
bullet("Création d'un tarball release.tar.gz")
bullet("Upload de l'artifact vers GitHub")

h3("9.2 Job 2 — deploy (dépend de test-and-build)")
bullet("Download de l'artifact")
bullet("Setup SSH agent avec la clé privée (secret SSH_PRIVATE_KEY)")
bullet("Ajout du VPS dans known_hosts (ssh-keyscan)")
bullet("scp du tarball et du script deploy.sh vers /tmp/ sur le VPS")
bullet("Exécution de deploy.sh à distance via SSH")
bullet("Smoke test : curl -fsS https://globalafricaplus.com")

h3("9.3 Environnement GitHub")
para("Un environment production a été configuré sur GitHub avec l'URL associée https://globalafricaplus.com, permettant d'ajouter optionnellement des reviewers manuels pour valider chaque déploiement.")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# 10. DEPLOY.SH
# ═══════════════════════════════════════════════════════════
h1("10. Script de déploiement atomique")

para("Le script deploy/deploy.sh exécuté sur le VPS effectue les opérations suivantes dans cet ordre strict :")

steps = [
    ("1. Extraction", "Décompresse le tarball dans releases/TIMESTAMP_sha8/"),
    ("2. Symlinks partagés", ".env, storage/ et bootstrap/cache/ pointent vers shared/"),
    ("3. Permissions", "chown deploy:www-data, chmod 2775 pour les dossiers, 0664 pour les fichiers"),
    ("4. APP_KEY", "Généré si absent du .env partagé (php artisan key:generate)"),
    ("5. Cache Laravel", "config:cache, route:cache, view:cache, event:cache"),
    ("6. Storage link", "php artisan storage:link (public/storage → storage/app/public)"),
    ("7. Migrations", "php artisan migrate --force --no-interaction"),
    ("8. SWITCH ATOMIQUE", "ln -sfn + mv -Tf sur current (zero downtime)"),
    ("9. Reload services", "sudo systemctl reload php8.2-fpm + supervisorctl restart workers"),
    ("10. Cleanup", "Conserve les 5 dernières releases, supprime les plus anciennes"),
]
for k, v in steps:
    p = doc.add_paragraph()
    r1 = p.add_run(k + " — ")
    r1.bold = True
    r1.font.color.rgb = GOLD
    p.add_run(v)

h3("Correctif appliqué pendant le déploiement")
para("Lors du premier run, deploy.sh a échoué sur un chmod de worker.log créé par supervisor en root. Le fix suivant a été poussé :")
code("""# Avant
chmod -R ug+rwX \"${SHARED_PATH}/storage\" \"${SHARED_PATH}/bootstrap/cache\"

# Après (commit f3a0887)
chmod -R ug+rwX \"${SHARED_PATH}/storage\" \"${SHARED_PATH}/bootstrap/cache\" 2>/dev/null || true""")

# ═══════════════════════════════════════════════════════════
# 11. EXÉCUTION
# ═══════════════════════════════════════════════════════════
h1("11. Exécution des déploiements")

h3("11.1 Run #1 — commit 837aed7 (ÉCHEC)")
table_2col([
    ("Résultat", "Échec à l'étape Execute remote deployment"),
    ("Cause", "worker.log appartenait à root → chmod refusé"),
    ("Correctif", "Commit f3a0887 (deploy.sh tolère les échecs chmod)"),
], header=("Item", "Détail"))

h3("11.2 Run #2 — commit f3a0887 (PARTIEL)")
para("Toutes les étapes critiques sont passées :")
bullet("APP_KEY généré")
bullet("Config / routes / views / events cachés")
bullet("Storage link créé")
bullet("14 migrations Laravel exécutées avec succès")
bullet("Switch atomique du symlink current effectué")
bullet("Site public : HTTP 200 OK")
para("Seul échec : sudo systemctl reload php8.2-fpm (sudoers NOPASSWD à configurer).", italic=True)

h3("11.3 Migrations exécutées")
migrations = [
    "0001_01_01_000000_create_users_table",
    "0001_01_01_000001_create_cache_table",
    "0001_01_01_000002_create_jobs_table",
    "2026_04_08_100000_extend_users_and_create_africaplus_tables",
    "2026_04_08_110000_create_role_profiles_table",
    "2026_04_09_100000_enrich_projects_sectors",
    "2026_04_09_120000_create_diaspora_module_tables",
    "2026_04_09_140000_enrich_mentorship_module",
    "2026_04_10_100000_enhance_government_module",
    "2026_04_10_200000_create_job_applications_table",
    "2026_04_11_100000_create_formalization_module",
    "2026_04_11_200000_create_subscription_module",
    "2026_04_12_100000_create_kyc_module",
    "2026_04_12_200000_create_advertising_module",
]
for m in migrations:
    p = doc.add_paragraph(style='List Number')
    r = p.add_run(m)
    r.font.name = 'Consolas'
    r.font.size = Pt(9.5)

h3("11.4 Run #3 — commit 00d5d84 (SUCCÈS 100%)")
para("Tous les jobs et étapes au vert :")
code("""✓ Tests & build frontend (1m13s)
  ✓ Setup PHP 8.2, Install Composer, Lint, Node 20, npm ci
  ✓ Build frontend (Vite)
  ✓ Package release + Upload artifact

✓ Deploy to production VPS (29s)
  ✓ Download artifact
  ✓ Setup SSH key + known_hosts
  ✓ Upload tarball + deploy.sh
  ✓ Execute remote deployment      ← sudoers OK
  ✓ Smoke test                     ← https://globalafricaplus.com 200 OK""")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# 12. VALIDATION
# ═══════════════════════════════════════════════════════════
h1("12. Validation finale et statut")

h3("12.1 Checklist complète")
items = [
    ("VPS Hostinger provisionné (LEMP + Redis + Supervisor)", "✅"),
    ("Utilisateur deploy avec clé SSH ed25519", "✅"),
    ("Sudoers NOPASSWD (reload php-fpm, restart workers)", "✅"),
    ("DNS A @ et A www → IP VPS", "✅"),
    ("SSL Let's Encrypt + renouvellement automatique", "✅"),
    ("HSTS + headers de sécurité Nginx", "✅"),
    (".env production complet (DB + Mail log + IDNorm sandbox)", "✅"),
    ("Base MySQL globalafricaplus + 14 migrations", "✅"),
    ("Pipeline GitHub Actions CI/CD", "✅"),
    ("Workers Supervisor (queue:work)", "✅"),
    ("Site public https://globalafricaplus.com", "✅ LIVE"),
]
t = doc.add_table(rows=1, cols=2)
t.style = 'Light Grid Accent 1'
hdr = t.rows[0].cells
hdr[0].text = "Élément"
hdr[1].text = "Statut"
for c in hdr:
    set_cell_bg(c, '1A1A1A')
    for p in c.paragraphs:
        for r in p.runs:
            r.font.color.rgb = WHITE
            r.bold = True
for lab, st in items:
    row = t.add_row().cells
    row[0].text = lab
    row[1].text = st
    for p in row[1].paragraphs:
        for r in p.runs:
            r.bold = True
            r.font.color.rgb = GREEN

h3("12.2 Réponse HTTP du site en production")
code("""$ curl -I https://globalafricaplus.com
HTTP/1.1 200 OK
Server: nginx/1.24.0 (Ubuntu)
Content-Type: text/html; charset=utf-8
Cache-Control: no-cache, private
Strict-Transport-Security: max-age=31536000; includeSubDomains
X-Frame-Options: SAMEORIGIN
X-Content-Type-Options: nosniff

$ curl -s https://globalafricaplus.com | grep title
<title>GlobalAfrica+ — Connecter l'Afrique et sa Diaspora</title>""")

# ═══════════════════════════════════════════════════════════
# 13. POST-DEPLOY
# ═══════════════════════════════════════════════════════════
h1("13. Procédures post-déploiement")

h3("13.1 Déploiement continu")
para("À partir de maintenant, chaque push sur main déclenche automatiquement :")
code("""git add .
git commit -m "feat: nouvelle fonctionnalité"
git push origin main

→ Workflow GitHub Actions (~2 min total)
→ Site mis à jour sans interruption""")

h3("13.2 Rollback rapide")
code("""cd /var/www/globalafricaplus.com/releases
ls -1tr                                    # liste des releases conservées
ln -sfn releases/PREVIOUS ../current       # pointe vers la précédente
sudo systemctl reload php8.2-fpm""")

h3("13.3 Consultation des logs")
code("""tail -f /var/www/globalafricaplus.com/shared/storage/logs/laravel.log
tail -f /var/log/nginx/globalafricaplus.error.log
sudo journalctl -u php8.2-fpm -f
sudo supervisorctl status""")

h3("13.4 Backup base de données (cron recommandé)")
code("""# Crontab root — backup quotidien à 3h
0 3 * * * mysqldump -u gapuser -p'MDP' globalafricaplus | gzip > /backup/gap_$(date +\\%F).sql.gz""")

h3("13.5 Étapes produit à venir")
bullet("Remplacer MAIL_MAILER=log par SMTP Hostinger (no-reply@globalafricaplus.com)")
bullet("Remplacer les credentials IDNorm sandbox par les credentials production")
bullet("Configurer un cron quotidien de backup DB (mysqldump → stockage externe)")
bullet("Monitoring externe : UptimeRobot / BetterStack sur https://globalafricaplus.com")
bullet("Taggage sémantique des versions (git tag v1.0.0 && git push --tags)")

hr_line()

# ═══════════════════════════════════════════════════════════
# FOOTER
# ═══════════════════════════════════════════════════════════
ft = doc.add_paragraph()
ft.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = ft.add_run("Mise en production réalisée le 18 avril 2026")
r.font.size = Pt(10)
r.italic = True
r.font.color.rgb = GREY

ft2 = doc.add_paragraph()
ft2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = ft2.add_run("GlobalAfrica+ — Connecter l'Afrique et sa Diaspora")
r.font.size = Pt(10)
r.bold = True
r.font.color.rgb = GOLD

# ═══════════════════════════════════════════════════════════
# SAVE
# ═══════════════════════════════════════════════════════════
out = r"C:\xampp\htdocs\globalafrica+\deploy\Rapport_Deploiement_GlobalAfricaPlus.docx"
doc.save(out)
print(f"OK - Document saved: {out}")
